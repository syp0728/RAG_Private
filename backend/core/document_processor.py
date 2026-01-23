"""
=============================================================================
DocumentProcessor - 문서 파싱 및 표 추출 모듈
=============================================================================

이 모듈은 PDF, DOCX, Excel 등 다양한 문서에서 텍스트와 표를 추출합니다.
특히 표(Table) 처리를 위해 5가지 방법을 지원합니다:

┌─────────────────────────────────────────────────────────────────────────────┐
│                        📊 표(Table) 처리 방법 5가지                          │
├─────────────────────────────────────────────────────────────────────────────┤
│                                                                             │
│  1️⃣ pdfplumber 텍스트 표 추출                                               │
│     - 방식: PDF 내부의 텍스트 기반 표 구조를 직접 파싱                        │
│     - 장점: 빠르고 정확함, 텍스트 기반 PDF에 최적                            │
│     - 단점: 이미지로 된 표는 인식 불가                                       │
│     - 메서드: _process_pdf_with_pdfplumber(), page.extract_tables()         │
│                                                                             │
│  2️⃣ 이미지 OCR (EasyOCR)                                                    │
│     - 방식: 페이지를 이미지로 변환 후 OCR로 텍스트 추출                       │
│     - 장점: 스캔 문서, 이미지 PDF 처리 가능                                  │
│     - 단점: 표 구조 인식 어려움, 속도 느림                                   │
│     - 메서드: _extract_image_tables_with_ocr(), reader.readtext()           │
│                                                                             │
│  3️⃣ OpenCV 표 선 감지                                                       │
│     - 방식: 이미지에서 수평/수직 선을 감지하여 셀 영역 분리                   │
│     - 장점: 선이 있는 표 정확하게 인식, 셀 단위 OCR 가능                     │
│     - 단점: 선이 없는 표는 인식 불가                                        │
│     - 메서드: _detect_table_cells_opencv(), cv2.morphologyEx()              │
│                                                                             │
│  4️⃣ EasyOCR 좌표 기반 추론                                                  │
│     - 방식: OCR 결과의 좌표(bbox)를 분석하여 행/열 구조 추론                  │
│     - 장점: 선이 없는 표도 처리 가능                                        │
│     - 단점: 복잡한 표 구조에서 오류 발생 가능                                │
│     - 메서드: _extract_image_tables_with_ocr() 내 좌표 그룹화 로직           │
│                                                                             │
│  5️⃣ Column-first Contextual Table Parsing                                   │
│     - 방식: 열 단위로 순회하며 병합 셀 채우기 + 계층 구조 텍스트 생성         │
│     - 장점: 병합 셀 처리, LLM이 이해하기 쉬운 계층형 출력                     │
│     - 단점: 단순한 표에는 오버헤드                                          │
│     - 메서드: _pdfplumber_table_to_markdown(), _excel_table_to_markdown()   │
│                                                                             │
├─────────────────────────────────────────────────────────────────────────────┤
│                            🔄 처리 우선순위                                   │
├─────────────────────────────────────────────────────────────────────────────┤
│                                                                             │
│  PDF 처리 시:                                                               │
│    1. OpenCV 표 선 감지 (우선) → 성공 시 셀별 OCR                            │
│    2. pdfplumber 텍스트 표 추출 (폴백)                                       │
│    3. 일반 텍스트 추출                                                      │
│                                                                             │
│  Excel 처리 시:                                                             │
│    1. openpyxl/xlrd로 셀 데이터 + 병합 셀 정보 추출                          │
│    2. Column-first Forward Fill 적용                                        │
│    3. 계층형 텍스트 + Markdown 테이블 생성                                   │
│                                                                             │
└─────────────────────────────────────────────────────────────────────────────┘

사용법:
    processor = DocumentProcessor()
    chunks = processor.extract_text_with_layout(file_path)
    
    # chunks 구조:
    # [
    #     {"text": "...", "page": 1, "type": "text", "metadata": {...}},
    #     {"text": "[계층형 표 데이터]...", "page": 2, "type": "table", "metadata": {"has_table": True}},
    # ]
"""

import re
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document

# pdfplumber를 사용한 표 추출 (Python 3.14 호환)
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
    print("pdfplumber available: Table extraction enabled")
except ImportError:
    HAS_PDFPLUMBER = False
    print("Warning: pdfplumber not available, table extraction disabled")

# openpyxl을 사용한 엑셀 처리
try:
    import openpyxl
    HAS_OPENPYXL = True
    print("openpyxl available: Excel file support enabled")
except ImportError:
    HAS_OPENPYXL = False
    print("Warning: openpyxl not available, Excel file support disabled")

# xlrd를 사용한 .xls 파일 처리
try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False

# unstructured는 선택사항 (Python 3.13 미만에서만 작동)
try:
    from unstructured.partition.auto import partition
    from unstructured.chunking.title import chunk_by_title
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False
    print("Warning: unstructured not available, using pdfplumber or PyPDF2 fallback")

# EasyOCR을 사용한 이미지 텍스트 추출
try:
    import easyocr
    HAS_EASYOCR = True
    print("EasyOCR available: Image text extraction enabled")
except ImportError:
    HAS_EASYOCR = False
    print("Warning: EasyOCR not available, image text extraction disabled")

# PIL을 사용한 이미지 처리
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    print("Warning: Pillow not available, image processing disabled")

# OpenCV를 사용한 표 선 감지
try:
    import cv2
    HAS_CV2 = True
    print("OpenCV available: Advanced table detection enabled")
except ImportError:
    HAS_CV2 = False
    print("Warning: OpenCV not available, advanced table detection disabled")

# pdf2image를 사용한 PDF 이미지 변환
try:
    from pdf2image import convert_from_path
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False
    print("Warning: pdf2image not available, PDF OCR disabled")

class DocumentProcessor:
    """Layout-aware 문서 처리 클래스"""
    
    def __init__(self):
        self.chunk_size = 1000
        self.chunk_overlap = 200
        self.ocr_reader = None  # Lazy loading for EasyOCR
    
    def _get_ocr_reader(self):
        """OCR 리더 초기화 (지연 로딩)"""
        if self.ocr_reader is None and HAS_EASYOCR:
            print("[DocumentProcessor] EasyOCR 모델 로딩 중... (처음 실행 시 시간이 소요됩니다)")
            self.ocr_reader = easyocr.Reader(['ko', 'en'], gpu=False)  # 한국어 + 영어 지원
            print("[DocumentProcessor] EasyOCR 모델 로딩 완료")
        return self.ocr_reader
    
    def extract_text_with_layout(self, file_path: Path) -> List[Dict]:
        """
        문서에서 텍스트, 표, 이미지를 추출하여 구조화된 청크 리스트 반환
        각 청크는 페이지 번호와 메타데이터를 포함
        """
        file_ext = file_path.suffix.lower()
        
        if file_ext == ".pdf":
            return self._process_pdf(file_path)
        elif file_ext == ".docx":
            return self._process_docx(file_path)
        elif file_ext in [".txt", ".md"]:
            return self._process_text(file_path)
        elif file_ext in [".xlsx", ".xls"]:
            return self._process_excel(file_path)
        elif file_ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"]:
            return self._process_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _process_pdf(self, file_path: Path) -> List[Dict]:
        """PDF 처리 (스마트 표 감지 + 자동 추출)"""
        chunks = []
        
        # ====== 스마트 표 감지 모드 ======
        # pdfplumber가 있으면 먼저 표가 있는지 확인하고,
        # 표가 있으면 pdfplumber로, 없으면 PyPDF2로 처리
        
        if HAS_PDFPLUMBER:
            try:
                # 1단계: 표 존재 여부 빠르게 확인
                has_tables = False
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages[:5]:  # 처음 5페이지만 검사 (속도 최적화)
                        tables = page.extract_tables()
                        if tables and any(t for t in tables if t and len(t) > 1):
                            has_tables = True
                            break
                
                if has_tables:
                    # 표가 감지됨 → pdfplumber로 표+텍스트 모두 추출
                    print(f"[DocumentProcessor] 표 감지됨! pdfplumber로 표 추출 모드 활성화")
                    self._process_pdf_with_pdfplumber(file_path, chunks)
                    
                    # 표 청크 수 카운트
                    table_count = sum(1 for c in chunks if c.get("type") == "table")
                    text_count = sum(1 for c in chunks if c.get("type") == "text")
                    print(f"[DocumentProcessor] pdfplumber 처리 완료: 텍스트 {text_count}개, 표 {table_count}개")
                else:
                    # 표 없음 → PyPDF2로 빠르게 텍스트만 추출
                    print(f"[DocumentProcessor] 표 없음, PyPDF2로 텍스트 추출")
                    self._process_pdf_with_pypdf2(file_path, chunks)
                    print(f"[DocumentProcessor] PyPDF2 처리 완료: {len(chunks)} 청크")
                    
            except Exception as e:
                print(f"[DocumentProcessor] pdfplumber 오류: {e}, PyPDF2로 폴백")
                chunks = []
                self._process_pdf_with_pypdf2(file_path, chunks)
        
        # pdfplumber가 없으면 unstructured 시도
        elif HAS_UNSTRUCTURED:
            try:
                # Unstructured를 사용한 구조화된 추출
                elements = partition(filename=str(file_path), strategy="hi_res")
                
                page_num = 1
                current_text = ""
                
                for element in elements:
                    element_text = str(element)
                    
                    # 페이지 번호 감지
                    if hasattr(element, 'metadata') and hasattr(element.metadata, 'page_number'):
                        new_page = element.metadata.page_number
                        if new_page != page_num:
                            # 페이지 변경 시 현재 청크 저장
                            if current_text.strip():
                                chunks.append({
                                    "text": current_text.strip(),
                                    "page": page_num,
                                    "type": "text"
                                })
                            current_text = ""
                            page_num = new_page
                    
                    # 표(table) 처리
                    if element.category == "Table":
                        # 표를 구조화된 텍스트 형식으로 변환
                        table_text = self._table_to_markdown(element_text)
                        chunks.append({
                            "text": f"\n\n[표 시작]\n{table_text}\n[표 끝]\n\n",
                            "page": page_num,
                            "type": "table"
                        })
                    else:
                        current_text += element_text + "\n\n"
                
                # 마지막 청크 저장
                if current_text.strip():
                    chunks.append({
                        "text": current_text.strip(),
                        "page": page_num,
                        "type": "text"
                    })
            
                print(f"[DocumentProcessor] unstructured로 PDF 처리 완료: {len(chunks)} 청크")
            
            except Exception as e:
                # Fallback: pdfplumber 사용
                print(f"[DocumentProcessor] Unstructured 처리 실패: {e}")
                chunks = []
                if HAS_PDFPLUMBER:
                    try:
                        self._process_pdf_with_pdfplumber(file_path, chunks)
                        print(f"[DocumentProcessor] pdfplumber로 PDF 처리 완료: {len(chunks)} 청크")
                    except Exception as e2:
                        print(f"[DocumentProcessor] pdfplumber 처리 실패: {e2}")
                        chunks = []
                        self._process_pdf_with_pypdf2(file_path, chunks)
                else:
                    self._process_pdf_with_pypdf2(file_path, chunks)
        
        # 2순위: pdfplumber 사용
        elif HAS_PDFPLUMBER:
            try:
                self._process_pdf_with_pdfplumber(file_path, chunks)
                print(f"[DocumentProcessor] pdfplumber로 PDF 처리 완료: {len(chunks)} 청크")
            except Exception as e:
                print(f"[DocumentProcessor] pdfplumber 처리 실패: {e}")
                chunks = []
                self._process_pdf_with_pypdf2(file_path, chunks)
        
        else:
            # 둘 다 없으면 PyPDF2 사용
            self._process_pdf_with_pypdf2(file_path, chunks)
        
        # 청크 분할
        return self._chunk_documents(chunks)
    
    def _process_pdf_with_pdfplumber(self, file_path: Path, chunks: List[Dict]):
        """
        ========================================================================
        [방법 1] pdfplumber를 사용한 PDF 처리
        ========================================================================
        
        처리 순서:
        1. OpenCV 표 선 감지 (우선) - _detect_table_cells_opencv()
        2. pdfplumber 텍스트 표 추출 (폴백) - page.extract_tables()
        3. 일반 텍스트 추출 - page.extract_text()
        
        장점: 텍스트 기반 PDF에서 빠르고 정확한 표 추출
        단점: 이미지로 된 표는 인식 불가 (OpenCV/OCR로 대체)
        """
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                table_bboxes = []  # 표 영역 좌표 저장 (텍스트 중복 방지용)
                tables_found = 0
                
                # ====== 1단계: OpenCV 기반 표 감지 (우선) ======
                if HAS_CV2 and HAS_EASYOCR and HAS_PIL:
                    print(f"[OpenCV] 페이지 {page_num}: OpenCV 표 감지 시도 (우선)")
                    image_tables = self._extract_image_tables_with_ocr(page, page_num)
                    if image_tables:
                        chunks.extend(image_tables)
                        tables_found += len(image_tables)
                        print(f"[OpenCV] 페이지 {page_num}: OpenCV로 표 {len(image_tables)}개 추출 성공!")
                    else:
                        print(f"[OpenCV] 페이지 {page_num}: OpenCV 표 감지 실패 -> pdfplumber로 전환")
                
                # ====== 2단계: pdfplumber 텍스트 기반 표 추출 (OpenCV 실패 시) ======
                if tables_found == 0:
                    tables = page.extract_tables()
                    
                    if tables:
                        for table_idx, table in enumerate(tables):
                            if table and len(table) > 1:  # 최소 2행 이상
                                # 표를 Markdown으로 변환 (Cell Merging + Fill-down 적용)
                                table_text = self._pdfplumber_table_to_markdown(table)
                                if table_text.strip():
                                    chunks.append({
                                        "text": f"\n\n[표 {table_idx + 1} 시작]\n{table_text}\n[표 {table_idx + 1} 끝]\n\n",
                                        "page": page_num,
                                        "type": "table"
                                    })
                                    tables_found += 1
                                    
                                    # 표 미리보기 로그 (처음 3행만)
                                    preview_lines = table_text.split('\n')[:5]
                                    preview = '\n    '.join(preview_lines)
                                    print(f"[pdfplumber TABLE] 페이지 {page_num}, 표 {table_idx + 1} ({len(table)}행 x {len(table[0]) if table[0] else 0}열)")
                                    print(f"    {preview}")
                                    if len(table_text.split('\n')) > 5:
                                        print(f"    ... (총 {len(table)}행)")
                
                # ====== 3단계: 일반 텍스트 추출 ======
                text = page.extract_text()
                if text and text.strip():
                    chunks.append({
                        "text": text.strip(),
                        "page": page_num,
                        "type": "text"
                    })
    
    def _pdfplumber_table_to_markdown(self, table: List[List]) -> str:
        """
        ========================================================================
        [방법 5] Column-first Contextual Table Parsing
        ========================================================================
        
        pdfplumber로 추출한 표를 계층형 텍스트 + Markdown으로 변환
        
        핵심 로직:
        1. Column-wise Forward Fill: 열 단위로 순회하며 빈 셀을 위쪽 값으로 채움
        2. Hierarchical Text Construction: "열1 > 열2 > 열3: 값" 형태로 계층 구조 표현
        3. Empty Cell Intelligence: 병합된 셀 정보를 참조하여 데이터 채우기
        
        출력 형식:
            [계층형 표 데이터]
              - 대분류 > 중분류 > 항목 >> 금액: 1000원
            
            [표 원본 (Markdown)]
            | 열1 | 열2 | 열3 |
            | --- | --- | --- |
            | ... | ... | ... |
        
        장점: LLM이 표 조각을 받아도 상위 계층 구조를 잃지 않음
        """
        if not table or len(table) < 1:
            return ""
        
        # ====== 1단계: 셀 정리 및 빈 행 제거 ======
        cleaned_table = []
        for row in table:
            if row and any(cell for cell in row):
                # None, 줄바꿈, 공백 정리
                cleaned_row = []
                for cell in row:
                    if cell is None:
                        cleaned_row.append("")
                    else:
                        # 줄바꿈을 공백으로 변환, 연속 공백 제거
                        cell_text = str(cell).replace("\n", " ").strip()
                        cell_text = " ".join(cell_text.split())
                        cleaned_row.append(cell_text)
                cleaned_table.append(cleaned_row)
        
        if not cleaned_table:
            return ""
        
        # ====== 2단계: 열 수 맞추기 ======
        max_cols = max(len(row) for row in cleaned_table)
        for row in cleaned_table:
            while len(row) < max_cols:
                row.append("")
        
        # ====== 3단계: Cell Merging (Fill-down) 처리 ======
        # 모든 열에 대해 빈 셀은 위쪽 값으로 채움
        for col_idx in range(max_cols):
            last_value = ""
            for row_idx in range(len(cleaned_table)):
                cell_value = cleaned_table[row_idx][col_idx]
                if cell_value:
                    last_value = cell_value
                elif last_value and row_idx > 0:
                    # 빈 셀: 위쪽 값으로 채움 (Fill-down)
                    cleaned_table[row_idx][col_idx] = last_value
        
        # ====== 4단계: 빈 열 제거 ======
        cols_to_keep = []
        for col_idx in range(max_cols):
            if any(row[col_idx].strip() for row in cleaned_table):
                cols_to_keep.append(col_idx)
        
        if cols_to_keep:
            cleaned_table = [[row[col_idx] for col_idx in cols_to_keep] for row in cleaned_table]
            max_cols = len(cols_to_keep)
        
        if max_cols == 0:
            return ""
        
        # ====== 5단계: 계층형 텍스트 + Markdown Table 생성 ======
        
        # 헤더 행 (첫 번째 행)
        headers = cleaned_table[0]
        # 빈 헤더는 "열1", "열2"... 로 채움
        for i, h in enumerate(headers):
            if not h.strip():
                headers[i] = f"열{i+1}"
        
        output_lines = []
        
        # ====== 계층형 텍스트 생성 (Hierarchical Text Construction) ======
        # 각 행을 "열1 > 열2 > 열3: 값" 형태로 변환
        output_lines.append("[계층형 표 데이터]")
        
        for row_idx, row in enumerate(cleaned_table[1:], 1):
            # 계층 구조 파악: 왼쪽부터 카테고리, 오른쪽이 값
            hierarchy_parts = []
            value_parts = []
            
            for col_idx, cell in enumerate(row):
                if not cell.strip():
                    continue
                    
                header = headers[col_idx] if col_idx < len(headers) else f"열{col_idx+1}"
                
                # 마지막 열 또는 숫자/금액이 포함된 열은 값으로 처리
                is_value = (col_idx >= len(row) - 2) or \
                           any(c.isdigit() for c in cell) or \
                           any(unit in cell for unit in ['원', '%', '개', '건', '명', '일'])
                
                if is_value and hierarchy_parts:
                    value_parts.append(f"{header}: {cell}")
                else:
                    hierarchy_parts.append(cell)
            
            # 계층 구조 문장 생성
            if hierarchy_parts or value_parts:
                if hierarchy_parts and value_parts:
                    hierarchy_str = " > ".join(hierarchy_parts)
                    value_str = ", ".join(value_parts)
                    output_lines.append(f"  - {hierarchy_str} >> {value_str}")
                elif hierarchy_parts:
                    output_lines.append(f"  - " + " > ".join(hierarchy_parts))
                elif value_parts:
                    output_lines.append(f"  - " + ", ".join(value_parts))
        
        output_lines.append("")
        
        # ====== Markdown Table도 함께 생성 (참조용) ======
        output_lines.append("[표 원본 (Markdown)]")
        output_lines.append("| " + " | ".join(headers) + " |")
        output_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        
        for row in cleaned_table[1:]:
            escaped_row = [cell.replace("|", "\\|") for cell in row]
            output_lines.append("| " + " | ".join(escaped_row) + " |")
        
        return "\n".join(output_lines)
    
    def _detect_table_cells_opencv(self, image_np, line_min_width=15):
        """
        ========================================================================
        [방법 3] OpenCV 표 선 감지
        ========================================================================
        
        이미지에서 수평/수직 선을 감지하여 표의 셀 영역을 분리합니다.
        
        처리 단계:
        1. 그레이스케일 변환: cv2.cvtColor(BGR2GRAY)
        2. 이진화: cv2.threshold() - 선을 흑/백으로 구분
        3. 수평선 감지: morphologyEx(MORPH_OPEN, kernal_h)
        4. 수직선 감지: morphologyEx(MORPH_OPEN, kernal_v)
        5. 선 결합: img_bin_h | img_bin_v
        6. 팽창: cv2.dilate() - 끊어진 선 연결
        7. 셀 영역 추출: cv2.connectedComponentsWithStats()
        
        Args:
            image_np: numpy array 형태의 이미지
            line_min_width: 선으로 인식할 최소 픽셀 크기 (기본 15px)
            
        Returns:
            list of (x, y, w, h) 튜플 - 각 셀의 위치와 크기
        
        장점: 선이 있는 표를 정확하게 셀 단위로 분리
        단점: 선이 없는 표는 인식 불가
        """
        import numpy as np
        
        if not HAS_CV2:
            return []
        
        try:
            # 1. 그레이스케일 변환
            if len(image_np.shape) == 3:
                gray_scale = cv2.cvtColor(image_np, cv2.COLOR_BGR2GRAY)
            else:
                gray_scale = image_np
            
            # 2. 이진화 (Threshold)
            _, img_bin = cv2.threshold(gray_scale, 150, 255, cv2.THRESH_BINARY)
            img_bin = ~img_bin  # 반전 (선이 흰색이 되도록)
            
            # 3. 수평/수직 커널 생성
            kernal_h = np.ones((1, line_min_width), np.uint8)  # 수평선 감지
            kernal_v = np.ones((line_min_width, 1), np.uint8)  # 수직선 감지
            
            # 4. 모폴로지 연산으로 선 감지
            img_bin_h = cv2.morphologyEx(img_bin, cv2.MORPH_OPEN, kernal_h)
            img_bin_v = cv2.morphologyEx(img_bin, cv2.MORPH_OPEN, kernal_v)
            
            # 5. 수평선 + 수직선 결합
            img_bin_final = img_bin_h | img_bin_v
            
            # 6. 팽창으로 선 연결
            final_kernel = np.ones((3, 3), np.uint8)
            img_bin_final = cv2.dilate(img_bin_final, final_kernel, iterations=1)
            
            # 7. 연결된 컴포넌트 분석 (셀 영역 감지)
            _, labels, stats, _ = cv2.connectedComponentsWithStats(
                img_bin_final, connectivity=8, ltype=cv2.CV_32S
            )
            
            # 8. 셀 영역 추출 (배경 제외, stats[0]은 배경)
            cells = []
            for i in range(2, len(stats)):  # 0: 배경, 1: 전체 표 테두리, 2+: 셀들
                x, y, w, h, area = stats[i]
                
                # 너무 작거나 너무 큰 영역 제외
                if w > 20 and h > 10 and area > 200 and area < (image_np.shape[0] * image_np.shape[1] * 0.5):
                    cells.append((x, y, w, h))
            
            print(f"[OpenCV] 표 셀 {len(cells)}개 감지됨")
            return cells
            
        except Exception as e:
            print(f"[OpenCV] 표 셀 감지 오류: {e}")
            return []
    
    def _ocr_table_cells(self, image_np, cells, page_num: int) -> List[Dict]:
        """
        ========================================================================
        [방법 3 보조] OpenCV로 감지된 셀에서 텍스트 추출
        ========================================================================
        
        _detect_table_cells_opencv()에서 감지된 셀 영역을 크롭하여
        각 셀별로 EasyOCR을 수행하고, 좌표 기반으로 행/열을 재구성합니다.
        
        처리 단계:
        1. 각 셀 영역 이미지 크롭
        2. EasyOCR로 텍스트 추출
        3. Y좌표로 행 그룹화 (허용 오차 20px)
        4. X좌표로 열 정렬
        5. Markdown 테이블로 변환
        
        Args:
            image_np: 원본 이미지 (numpy array)
            cells: (x, y, w, h) 튜플 리스트 - OpenCV에서 감지된 셀 좌표
            page_num: 페이지 번호
            
        Returns:
            청크 리스트 [{"text": "...", "page": N, "type": "table"}]
        """
        chunks = []
        
        if not cells or not HAS_EASYOCR:
            return chunks
        
        try:
            reader = self._get_ocr_reader()
            if not reader:
                return chunks
            
            import numpy as np
            
            # 각 셀에서 OCR 수행
            cell_data = []
            for (x, y, w, h) in cells:
                # 셀 영역 크롭 (약간의 패딩 추가)
                pad = 2
                y1, y2 = max(0, y + pad), min(image_np.shape[0], y + h - pad)
                x1, x2 = max(0, x + pad), min(image_np.shape[1], x + w - pad)
                
                cell_img = image_np[y1:y2, x1:x2]
                
                if cell_img.size == 0:
                    continue
                
                # OCR
                results = reader.readtext(cell_img, detail=0, paragraph=True)
                text = " ".join(results).strip() if results else ""
                
                # 셀 중심 좌표
                center_x = x + w // 2
                center_y = y + h // 2
                
                cell_data.append({
                    "text": text,
                    "x": center_x,
                    "y": center_y,
                    "w": w,
                    "h": h
                })
            
            if not cell_data:
                return chunks
            
            # Y좌표로 행 그룹화 (허용 오차 20픽셀)
            cell_data.sort(key=lambda c: c["y"])
            rows = []
            current_row = [cell_data[0]]
            current_y = cell_data[0]["y"]
            
            for cell in cell_data[1:]:
                if abs(cell["y"] - current_y) <= 20:
                    current_row.append(cell)
                else:
                    rows.append(sorted(current_row, key=lambda c: c["x"]))
                    current_row = [cell]
                    current_y = cell["y"]
            rows.append(sorted(current_row, key=lambda c: c["x"]))
            
            # Markdown 테이블로 변환
            table_rows = [[cell["text"] for cell in row] for row in rows]
            
            if len(table_rows) >= 2:
                table_text = self._pdfplumber_table_to_markdown(table_rows)
                
                if table_text.strip():
                    chunks.append({
                        "text": f"\n\n[OpenCV 표 감지 - {len(rows)}행 x {len(rows[0]) if rows else 0}열]\n{table_text}\n[표 끝]\n\n",
                        "page": page_num,
                        "type": "table"
                    })
                    
                    # 로그 출력
                    preview_lines = table_text.split('\n')[:5]
                    preview = '\n    '.join(preview_lines)
                    print(f"[OpenCV TABLE] 페이지 {page_num}: {len(rows)}행 x {len(rows[0]) if rows else 0}열")
                    print(f"    {preview}")
            
            return chunks
            
        except Exception as e:
            print(f"[OpenCV] OCR 표 추출 오류: {e}")
            return chunks
    
    def _extract_image_tables_with_ocr(self, page, page_num: int) -> List[Dict]:
        """
        ========================================================================
        [방법 2, 3, 4 통합] 이미지 기반 표 추출 (OCR)
        ========================================================================
        
        pdfplumber 페이지를 이미지로 변환한 후 표를 추출합니다.
        
        처리 순서 (우선순위대로):
        1. [방법 3] OpenCV 표 선 감지 → 셀별 OCR
           - _detect_table_cells_opencv()로 셀 영역 감지
           - _ocr_table_cells()로 각 셀 OCR
           
        2. [방법 4] EasyOCR 좌표 기반 추론 (OpenCV 실패 시)
           - reader.readtext()로 전체 OCR
           - bbox 좌표를 분석하여 행/열 구조 추론
           - Y좌표로 행 그룹화, X좌표로 열 정렬
           
        3. [방법 2] 단순 EasyOCR (구조화 실패 시)
           - 텍스트만 추출하여 반환
        
        장점: 스캔 문서, 이미지 PDF 처리 가능
        단점: 속도 느림, 복잡한 표에서 오류 가능
        """
        chunks = []
        
        if not HAS_EASYOCR or not HAS_PIL:
            return chunks
        
        try:
            # 페이지를 이미지로 변환
            page_image = page.to_image(resolution=150)
            pil_image = page_image.original
            
            # RGB 변환
            if pil_image.mode != 'RGB':
                pil_image = pil_image.convert('RGB')
            
            import numpy as np
            image_np = np.array(pil_image)
            
            # OpenCV로 표 셀 감지 시도
            if HAS_CV2:
                print(f"[OpenCV] 표 선 감지 시작 (이미지 크기: {image_np.shape[1]}x{image_np.shape[0]})")
                cells = self._detect_table_cells_opencv(image_np)
                if cells:
                    print(f"[OpenCV] 셀 {len(cells)}개 감지됨 -> EasyOCR로 셀 내용 추출 중...")
                    opencv_chunks = self._ocr_table_cells(image_np, cells, page_num)
                    if opencv_chunks:
                        print(f"[OpenCV] 표 추출 성공!")
                        return opencv_chunks
                    print(f"[OpenCV] 셀 OCR 실패 -> 기본 EasyOCR 방식으로 전환")
                else:
                    print(f"[OpenCV] 표 선 감지 실패 -> 기본 EasyOCR 방식으로 전환")
            else:
                print(f"[EasyOCR] OpenCV 비활성화 -> 좌표 기반 표 추론 방식 사용")
            
            # OpenCV 실패 시 기본 OCR 방식 사용
            print(f"[EasyOCR] 전체 페이지 OCR 수행 중...")
            reader = self._get_ocr_reader()
            if not reader:
                print(f"[EasyOCR] OCR 리더 생성 실패")
                return chunks
            
            # OCR 수행 (좌표 정보 포함)
            results = reader.readtext(image_np, detail=1, paragraph=False)
            print(f"[EasyOCR] {len(results) if results else 0}개 텍스트 영역 감지")
            
            if not results:
                return chunks
            
            # 좌표 기반으로 표 구조 추론
            # 1. Y좌표로 그룹화 (같은 행)
            rows = {}
            for (bbox, text, conf) in results:
                if conf < 0.3:  # 신뢰도 낮은 결과 제외
                    continue
                
                # bbox: [[x1,y1], [x2,y1], [x2,y2], [x1,y2]]
                y_center = (bbox[0][1] + bbox[2][1]) / 2
                x_center = (bbox[0][0] + bbox[2][0]) / 2
                
                # Y좌표를 20픽셀 단위로 그룹화 (같은 행으로 취급)
                row_key = int(y_center / 20) * 20
                
                if row_key not in rows:
                    rows[row_key] = []
                rows[row_key].append((x_center, text.strip()))
            
            if len(rows) < 2:  # 최소 2행 이상이어야 표로 인식
                print(f"[EasyOCR] 표로 인식하기엔 행이 부족 ({len(rows)}행)")
                return chunks
            
            # 2. 각 행을 X좌표로 정렬
            sorted_rows = []
            for row_key in sorted(rows.keys()):
                row_cells = sorted(rows[row_key], key=lambda x: x[0])
                row_texts = [cell[1] for cell in row_cells if cell[1]]
                if row_texts:
                    sorted_rows.append(row_texts)
            
            if len(sorted_rows) < 2:
                print(f"[EasyOCR] 정렬 후 행이 부족")
                return chunks
            
            # 3. Markdown 테이블로 변환
            print(f"[EasyOCR] 좌표 기반 표 구조화: {len(sorted_rows)}행 감지")
            table_text = self._pdfplumber_table_to_markdown(sorted_rows)
            
            if table_text.strip():
                chunks.append({
                    "text": f"\n\n[이미지 표 - EasyOCR 좌표 추론]\n{table_text}\n[이미지 표 끝]\n\n",
                    "page": page_num,
                    "type": "table"
                })
                print(f"[EasyOCR] 표 추출 성공! ({len(sorted_rows)}행)")
                # 표 미리보기
                preview_lines = table_text.split('\n')[:3]
                for line in preview_lines:
                    print(f"    {line}")
        
        except Exception as e:
            print(f"[DocumentProcessor] 이미지 표 OCR 오류 (페이지 {page_num}): {e}")
        
        return chunks
    
    def _process_pdf_with_pypdf2(self, file_path: Path, chunks: List[Dict]):
        """PyPDF2를 사용한 PDF 처리 (Fallback)"""
        with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        chunks.append({
                            "text": text.strip(),
                            "page": page_num,
                            "type": "text"
                        })
        
        # 청크 분할
        return self._chunk_documents(chunks)
    
    def _process_docx(self, file_path: Path) -> List[Dict]:
        """DOCX 처리"""
        chunks = []
        doc = Document(file_path)
        
        current_text = ""
        page_num = 1  # DOCX는 정확한 페이지 번호가 없으므로 1로 설정
        
        for para in doc.paragraphs:
            current_text += para.text + "\n\n"
        
        # 표 처리
        for table in doc.tables:
            table_text = self._docx_table_to_markdown(table)
            if current_text.strip():
                chunks.append({
                    "text": current_text.strip(),
                    "page": page_num,
                    "type": "text"
                })
                current_text = ""
            
            chunks.append({
                "text": f"\n\n[표 시작]\n{table_text}\n[표 끝]\n\n",
                "page": page_num,
                "type": "table"
            })
        
        if current_text.strip():
            chunks.append({
                "text": current_text.strip(),
                "page": page_num,
                "type": "text"
            })
        
        return self._chunk_documents(chunks)
    
    def _process_text(self, file_path: Path) -> List[Dict]:
        """텍스트 파일 처리"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        return self._chunk_documents([{
            "text": text,
            "page": 1,
            "type": "text"
        }])
    
    def _process_image(self, file_path: Path) -> List[Dict]:
        """이미지 파일에서 OCR로 텍스트 추출"""
        if not HAS_EASYOCR:
            raise ValueError("EasyOCR이 설치되지 않았습니다. 'pip install easyocr'로 설치하세요.")
        
        if not HAS_PIL:
            raise ValueError("Pillow가 설치되지 않았습니다. 'pip install Pillow'로 설치하세요.")
        
        chunks = []
        
        try:
            # OCR 리더 가져오기 (지연 로딩)
            reader = self._get_ocr_reader()
            
            # 이미지 로드
            image = Image.open(file_path)
            
            # RGB로 변환 (필요한 경우)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            print(f"[DocumentProcessor] OCR 처리 중: {file_path.name}")
            
            # OCR 수행
            result = reader.readtext(str(file_path), detail=0, paragraph=True)
            
            # 결과 텍스트 합치기
            extracted_text = "\n".join(result)
            
            if extracted_text.strip():
                chunks.append({
                    "text": extracted_text.strip(),
                    "page": 1,
                    "type": "ocr"
                })
                print(f"[DocumentProcessor] OCR 완료: {len(extracted_text)} 글자 추출")
            else:
                print(f"[DocumentProcessor] OCR 결과 없음: {file_path.name}")
                chunks.append({
                    "text": f"[이미지 파일: {file_path.name}] - 텍스트 추출 불가",
                    "page": 1,
                    "type": "ocr"
                })
        
        except Exception as e:
            print(f"[DocumentProcessor] OCR 오류: {e}")
            chunks.append({
                "text": f"[이미지 파일: {file_path.name}] - OCR 처리 오류: {str(e)}",
                "page": 1,
                "type": "ocr"
            })
        
        return self._chunk_documents(chunks)
    
    def _process_pdf_with_ocr(self, file_path: Path) -> List[Dict]:
        """PDF 페이지를 이미지로 변환 후 OCR 처리 (스캔 PDF용)"""
        if not HAS_PDF2IMAGE:
            raise ValueError("pdf2image가 설치되지 않았습니다. 'pip install pdf2image'로 설치하세요.")
        
        if not HAS_EASYOCR:
            raise ValueError("EasyOCR이 설치되지 않았습니다. 'pip install easyocr'로 설치하세요.")
        
        chunks = []
        
        try:
            # OCR 리더 가져오기
            reader = self._get_ocr_reader()
            
            print(f"[DocumentProcessor] PDF OCR 처리 중: {file_path.name}")
            
            # PDF를 이미지로 변환
            images = convert_from_path(str(file_path), dpi=200)
            
            for page_num, image in enumerate(images, 1):
                print(f"[DocumentProcessor] 페이지 {page_num}/{len(images)} OCR 처리 중...")
                
                # OCR 수행
                import numpy as np
                image_np = np.array(image)
                result = reader.readtext(image_np, detail=0, paragraph=True)
                
                # 결과 텍스트 합치기
                page_text = "\n".join(result)
                
                if page_text.strip():
                    chunks.append({
                        "text": page_text.strip(),
                        "page": page_num,
                        "type": "ocr"
                    })
            
            print(f"[DocumentProcessor] PDF OCR 완료: {len(chunks)} 페이지 처리")
        
        except Exception as e:
            print(f"[DocumentProcessor] PDF OCR 오류: {e}")
            # Fallback to PyPDF2
            chunks = []
            self._process_pdf_with_pypdf2(file_path, chunks)
        
        return self._chunk_documents(chunks)
    
    def _process_excel(self, file_path: Path) -> List[Dict]:
        """엑셀 파일 처리 (.xlsx, .xls)"""
        chunks = []
        file_ext = file_path.suffix.lower()
        
        if file_ext == ".xlsx" and HAS_OPENPYXL:
            # .xlsx 파일 처리
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_idx, sheet_name in enumerate(workbook.sheetnames, 1):
                sheet = workbook[sheet_name]
                
                # 병합 셀 정보 수집 (Merged Cells Info)
                merged_cells_map = {}
                for merged_range in sheet.merged_cells.ranges:
                    min_row, min_col = merged_range.min_row, merged_range.min_col
                    max_row, max_col = merged_range.max_row, merged_range.max_col
                    
                    # 병합 셀의 첫 번째 값 가져오기
                    first_cell_value = sheet.cell(min_row, min_col).value
                    first_cell_value = str(first_cell_value) if first_cell_value else ""
                    
                    # 병합된 모든 셀에 같은 값 매핑
                    for r in range(min_row, max_row + 1):
                        for c in range(min_col, max_col + 1):
                            merged_cells_map[(r, c)] = first_cell_value
                
                print(f"[Excel] 시트 '{sheet_name}': 병합 셀 {len(sheet.merged_cells.ranges)}개 감지")
                
                # 시트의 모든 데이터를 표로 추출 (병합 셀 정보 적용)
                table_data = []
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    processed_row = []
                    for col_idx, cell in enumerate(row, 1):
                        # 병합 셀 정보가 있으면 사용
                        if (row_idx, col_idx) in merged_cells_map:
                            processed_row.append(merged_cells_map[(row_idx, col_idx)])
                        elif cell is not None:
                            processed_row.append(str(cell))
                        else:
                            processed_row.append("")
                    
                    if any(cell for cell in processed_row):
                        table_data.append(processed_row)
                
                if table_data:
                    # 마크다운 테이블로 변환 (계층형 포함)
                    table_text = self._excel_table_to_markdown(table_data, sheet_name)
                    chunks.append({
                        "text": f"[시트: {sheet_name}]\n\n[표 시작]\n{table_text}\n[표 끝]",
                        "page": sheet_idx,
                        "type": "table"
                    })
                    
                    # 로그 출력
                    preview_lines = table_text.split('\n')[:5]
                    print(f"[Excel TABLE] 시트 '{sheet_name}': {len(table_data)}행")
                    for line in preview_lines:
                        print(f"    {line}")
            
            workbook.close()
            
        elif file_ext == ".xls" and HAS_XLRD:
            # .xls 파일 처리
            workbook = xlrd.open_workbook(file_path)
            
            for sheet_idx in range(workbook.nsheets):
                sheet = workbook.sheet_by_index(sheet_idx)
                sheet_name = sheet.name
                
                # 시트의 모든 데이터를 표로 추출
                table_data = []
                for row_idx in range(sheet.nrows):
                    row = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                    if any(cell for cell in row):
                        table_data.append(row)
                
                if table_data:
                    # 마크다운 테이블로 변환
                    table_text = self._excel_table_to_markdown(table_data, sheet_name)
                    chunks.append({
                        "text": f"[시트: {sheet_name}]\n\n[표 시작]\n{table_text}\n[표 끝]",
                        "page": sheet_idx + 1,
                        "type": "table"
                    })
        else:
            raise ValueError(f"Excel file processing not available for {file_ext}")
        
        print(f"[DocumentProcessor] Excel 처리 완료: {len(chunks)} 시트")
        return self._chunk_documents(chunks)
    
    def _excel_table_to_markdown(self, table_data: List[List], sheet_name: str = "") -> str:
        """엑셀 테이블 데이터를 계층형 텍스트 + 마크다운 형식으로 변환"""
        if not table_data:
            return ""
        
        # 최대 열 수 계산
        max_cols = max(len(row) for row in table_data)
        
        # 모든 행을 같은 열 수로 맞추기
        for row in table_data:
            while len(row) < max_cols:
                row.append("")
        
        # ====== Column-wise Forward Fill (열 단위 채우기) ======
        for col_idx in range(max_cols):
            last_value = ""
            for row_idx in range(len(table_data)):
                cell_value = table_data[row_idx][col_idx].strip()
                if cell_value:
                    last_value = cell_value
                elif last_value and row_idx > 0:
                    # 빈 셀: 위쪽 값으로 채움 (Fill-down)
                    table_data[row_idx][col_idx] = last_value
        
        # 빈 열 제거 (모든 행에서 빈 열)
        cols_to_keep = []
        for col_idx in range(max_cols):
            if any(row[col_idx].strip() for row in table_data):
                cols_to_keep.append(col_idx)
        
        if cols_to_keep:
            table_data = [[row[col_idx] for col_idx in cols_to_keep] for row in table_data]
            max_cols = len(cols_to_keep)
        
        output_lines = []
        
        # 헤더 행
        headers = table_data[0] if table_data else []
        for i, h in enumerate(headers):
            if not h.strip():
                headers[i] = f"열{i+1}"
        
        # ====== 계층형 텍스트 생성 (Hierarchical Text Construction) ======
        output_lines.append("[계층형 데이터]")
        
        for row in table_data[1:]:
            hierarchy_parts = []
            value_parts = []
            
            for col_idx, cell in enumerate(row):
                if not cell.strip():
                    continue
                    
                header = headers[col_idx] if col_idx < len(headers) else f"열{col_idx+1}"
                
                # 마지막 2열 또는 숫자/금액이 포함된 열은 값으로 처리
                is_value = (col_idx >= len(row) - 2) or \
                           any(c.isdigit() for c in cell) or \
                           any(unit in cell for unit in ['원', '%', '개', '건', '명', '일', '시간'])
                
                if is_value and hierarchy_parts:
                    value_parts.append(f"{header}: {cell}")
                else:
                    hierarchy_parts.append(cell)
            
            if hierarchy_parts or value_parts:
                if hierarchy_parts and value_parts:
                    hierarchy_str = " > ".join(hierarchy_parts)
                    value_str = ", ".join(value_parts)
                    output_lines.append(f"  - {hierarchy_str} >> {value_str}")
                elif hierarchy_parts:
                    output_lines.append(f"  - " + " > ".join(hierarchy_parts))
                elif value_parts:
                    output_lines.append(f"  - " + ", ".join(value_parts))
        
        output_lines.append("")
        
        # ====== Markdown Table (참조용) ======
        output_lines.append("[표 원본]")
        output_lines.append("| " + " | ".join(headers) + " |")
        output_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        
        for row in table_data[1:]:
            escaped_row = [cell.replace("|", "\\|") for cell in row]
            output_lines.append("| " + " | ".join(escaped_row) + " |")
        
        return "\n".join(output_lines)
    
    def _table_to_markdown(self, table_text: str) -> str:
        """표 텍스트를 마크다운 형식으로 변환"""
        lines = table_text.strip().split("\n")
        if not lines:
            return ""
        
        if len(lines) < 2:
            return table_text
        
        # 각 행을 셀로 분리
        all_rows = []
        max_cols = 0
        
        for line in lines:
            # 탭으로 분리 시도
            cells = [c.strip() for c in line.split("\t")]
            if len(cells) == 1:
                # 탭이 없으면 여러 공백으로 분리
                cells = [c.strip() for c in line.split("  ") if c.strip()]
            if not cells:
                cells = [line.strip()]
            
            all_rows.append(cells)
            max_cols = max(max_cols, len(cells))
        
        if max_cols == 0:
            return table_text
        
        # 모든 행을 같은 열 수로 맞추기
        for row in all_rows:
            while len(row) < max_cols:
                row.append("")
        
        # 마크다운 테이블 생성
        md_lines = []
        
        # 헤더 행
        md_lines.append("| " + " | ".join(all_rows[0]) + " |")
        
        # 구분선
        md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        
        # 데이터 행
        for row in all_rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(md_lines)
    
    def _docx_table_to_markdown(self, table) -> str:
        """DOCX 표를 마크다운 형식으로 변환"""
        md_lines = []
        max_cols = 0
        
        # 모든 행 수집 및 최대 열 수 계산
        all_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            all_rows.append(cells)
            max_cols = max(max_cols, len(cells))
        
        if not all_rows:
            return ""
        
        # 모든 행을 같은 열 수로 맞추기
        for row in all_rows:
            while len(row) < max_cols:
                row.append("")
        
        # 헤더 행
        md_lines.append("| " + " | ".join(all_rows[0]) + " |")
        
        # 구분선
        md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        
        # 데이터 행
        for row in all_rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(md_lines)
    
    def _chunk_documents(self, chunks: List[Dict]) -> List[Dict]:
        """문서를 지정된 크기로 청크 분할
        
        표(table) 청크는 분할하지 않고 온전히 하나의 청크로 유지
        has_table 메타데이터로 표 포함 여부 표시
        """
        final_chunks = []
        
        # 표 청크 최대 크기 (표는 일반 텍스트보다 크게 허용)
        TABLE_MAX_SIZE = self.chunk_size * 3  # 표는 3배 크기까지 허용
        
        for chunk in chunks:
            text = chunk["text"]
            page = chunk["page"]
            chunk_type = chunk["type"]
            is_table = chunk_type == "table"
            
            # ====== 표 청크: 분할하지 않고 온전히 유지 ======
            if is_table:
                # 표는 가급적 분할하지 않음 (매우 큰 표만 예외적으로 분할)
                if len(text) <= TABLE_MAX_SIZE:
                    final_chunks.append({
                        "text": text,
                        "page": page,
                        "type": chunk_type,
                        "metadata": {
                            "page": page, 
                            "type": chunk_type,
                            "has_table": True  # 표 포함 태그
                        }
                    })
                else:
                    # 매우 큰 표: 행 단위로 분할 (Markdown 테이블 구조 유지)
                    lines = text.split("\n")
                    header_lines = []
                    data_lines = []
                    
                    # 헤더와 구분선 추출
                    for i, line in enumerate(lines):
                        if i < 3 and (line.startswith("|") or "---" in line or "[표" in line):
                            header_lines.append(line)
                        else:
                            data_lines.append(line)
                    
                    header_text = "\n".join(header_lines)
                    current_chunk_lines = []
                    current_length = len(header_text)
                    
                    for line in data_lines:
                        if current_length + len(line) > self.chunk_size and current_chunk_lines:
                            # 현재 청크 저장 (헤더 포함)
                            chunk_text = header_text + "\n" + "\n".join(current_chunk_lines)
                            final_chunks.append({
                                "text": chunk_text,
                                "page": page,
                                "type": chunk_type,
                                "metadata": {
                                    "page": page, 
                                    "type": chunk_type,
                                    "has_table": True,
                                    "table_continued": True  # 이어지는 표임을 표시
                                }
                            })
                            current_chunk_lines = [line]
                            current_length = len(header_text) + len(line)
                        else:
                            current_chunk_lines.append(line)
                            current_length += len(line) + 1
                    
                    # 마지막 청크
                    if current_chunk_lines:
                        chunk_text = header_text + "\n" + "\n".join(current_chunk_lines)
                        final_chunks.append({
                            "text": chunk_text,
                            "page": page,
                            "type": chunk_type,
                            "metadata": {
                                "page": page, 
                                "type": chunk_type,
                                "has_table": True
                            }
                        })
            
            # ====== 일반 텍스트 청크 ======
            elif len(text) <= self.chunk_size:
                final_chunks.append({
                    "text": text,
                    "page": page,
                    "type": chunk_type,
                    "metadata": {
                        "page": page, 
                        "type": chunk_type,
                        "has_table": False
                    }
                })
            else:
                # 텍스트를 더 작은 청크로 분할
                words = text.split()
                current_chunk = []
                current_length = 0
                
                for word in words:
                    word_length = len(word) + 1
                    if current_length + word_length > self.chunk_size and current_chunk:
                        chunk_text = " ".join(current_chunk)
                        final_chunks.append({
                            "text": chunk_text,
                            "page": page,
                            "type": chunk_type,
                            "metadata": {
                                "page": page, 
                                "type": chunk_type,
                                "has_table": False
                            }
                        })
                        # Overlap을 위해 마지막 부분 유지
                        overlap_size = int(self.chunk_overlap / 10)
                        current_chunk = current_chunk[-overlap_size:] + [word]
                        current_length = sum(len(w) + 1 for w in current_chunk)
                    else:
                        current_chunk.append(word)
                        current_length += word_length
                
                # 마지막 청크
                if current_chunk:
                    chunk_text = " ".join(current_chunk)
                    final_chunks.append({
                        "text": chunk_text,
                        "page": page,
                        "type": chunk_type,
                        "metadata": {
                            "page": page, 
                            "type": chunk_type,
                            "has_table": False
                        }
                    })
        
        # 통계 로깅
        table_chunks = sum(1 for c in final_chunks if c["metadata"].get("has_table"))
        text_chunks = len(final_chunks) - table_chunks
        print(f"[DocumentProcessor] 청킹 완료: 텍스트 {text_chunks}개, 표 {table_chunks}개")
        
        return final_chunks

